from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = 'output/docx/Cuadro_analisis_Caso_Disneyland_Paris.docx'

def set_font(run, name='Calibri', size=10.5, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_width(cell, width_dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa)); tcW.set(qn('w:type'),'dxa')

def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc; tcPr=tc.get_or_add_tcPr(); margins=tcPr.first_child_found_in('w:tcMar')
    if margins is None: margins=OxmlElement('w:tcMar'); tcPr.append(margins)
    for side, value in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=margins.find(qn(f'w:{side}'))
        if node is None: node=OxmlElement(f'w:{side}'); margins.append(node)
        node.set(qn('w:w'),str(value)); node.set(qn('w:type'),'dxa')

def set_table_geometry(table, widths):
    tblPr=table._tbl.tblPr
    tblW=tblPr.first_child_found_in('w:tblW'); tblW.set(qn('w:w'),'9360'); tblW.set(qn('w:type'),'dxa')
    tblInd=OxmlElement('w:tblInd'); tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa'); tblPr.append(tblInd)
    layout=OxmlElement('w:tblLayout'); layout.set(qn('w:type'),'fixed'); tblPr.append(layout)
    grid=table._tbl.tblGrid
    for gc,w in zip(grid.gridCol_lst,widths): gc.set(qn('w:w'),str(w))
    for row in table.rows:
        for cell,w in zip(row.cells,widths): set_cell_width(cell,w); set_cell_margins(cell)

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); node=OxmlElement('w:tblHeader'); node.set(qn('w:val'),'true'); trPr.append(node)

def add_para(cell, text, bold=False, color=None, size=10.2):
    p=cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.12
    r=p.add_run(text); set_font(r,size=size,bold=bold,color=color)
    return p

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(0.85); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.85); sec.right_margin=Inches(0.85)
sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
styles=doc.styles
styles['Normal'].font.name='Calibri'; styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); styles['Normal'].font.size=Pt(10.5)

# Header/footer
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
hr=hp.add_run('Dirección Estratégica y Gestión de Personas'); set_font(hr,size=8.5,color=(89,89,89))
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=fp.add_run('Caso Disneyland Paris | Cuadro de análisis'); set_font(fr,size=8.5,color=(89,89,89))

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3)
r=p.add_run('CASO DISNEYLAND PARIS'); set_font(r,size=20,bold=True,color=(31,77,120))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(14)
r=p.add_run('Cuadro de análisis y respuestas para la sustentación'); set_font(r,size=11,color=(89,89,89))

call=doc.add_table(rows=1, cols=1); call.alignment=WD_TABLE_ALIGNMENT.CENTER
set_table_geometry(call,[9360]); c=call.cell(0,0); shade(c,'E8EEF5'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
add_para(c,'Idea central: Disney debe crecer con una lógica de cartera integrada: proteger el parque como motor de demanda y marca, mientras desarrolla Val d’Europe por etapas, con adaptación local y disciplina financiera.',bold=True,color=(31,58,95),size=10.5)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

questions=[
('¿Cómo describirían la gestión administrativa de Disney del parque temático de París, Fase III?',
 'Es una gestión estratégica de transformación y de portafolio. Integra expansión de experiencias (Marvel, Frozen y The Lion King), renovación hotelera, tecnología y entretenimiento nocturno. Su fortaleza es coordinar marca, diseño, operación y experiencia del visitante. Requiere gobernanza de proyectos: control de presupuesto, cronogramas, capacidad instalada, calidad de servicio y coordinación con autoridades locales.',
 'Gestionar por etapas y con tableros de control: demanda, rentabilidad de cada inversión, satisfacción del visitante, avances de obra y efectos sobre el destino.'),
('¿Cómo Disney debería administrar y priorizar la siguiente fase de crecimiento, tanto para el parque temático como para Val d’Europe?',
 'Debe aplicar una estrategia equilibrada, sin elegir de forma excluyente entre parque y Val d’Europe. El parque genera marca, visitantes y gasto turístico; Val d’Europe diversifica ingresos mediante comercio, oficinas, vivienda y hoteles.',
 'Priorizar primero inversiones del parque que sostengan diferenciación, demanda y retorno. Desarrollar Val d’Europe escalonadamente, condicionado a demanda comprobada. Crear una cartera conjunta Disney-autoridades que evalúe retorno, empleo, movilidad, sostenibilidad y riesgo.'),
('¿Qué tipo de incertidumbre afronta Disney al construir la Fase III y continuar el desarrollo de Val d’Europe?',
 'Afronta incertidumbre de mercado (demanda turística, gasto y preferencia por franquicias), macroeconómica y financiera (inflación, tasas, tipo de cambio y costos de obra), operativa y regulatoria (plazos, permisos, transporte y ambiente), y sociocultural-reputacional (aceptación local y coherencia de marca).',
 'Usar escenarios conservador, base y expansivo; establecer inversiones por hitos, reservas de contingencia, pilotos de demanda y alertas tempranas para corregir antes de comprometer más capital.'),
('¿Qué aprendizajes de la primera fase de Disneyland Paris debería considerar la compañía al desarrollar la Fase III y la región Val d’Europe?',
 'La experiencia inicial mostró que no se debe trasladar sin ajuste el modelo estadounidense. La adaptación cultural, gastronómica, arquitectónica y laboral fue decisiva. También evidenció que las previsiones de asistencia, ocupación y gasto pueden ser demasiado optimistas, y que deuda y sobrecostos ponen en riesgo el proyecto.',
 'Validar supuestos con datos locales; proteger liquidez; controlar costos y plazos; fortalecer relaciones con empleados, comunidades y autoridades; y gestionar parque, transporte, hoteles, comercio y vivienda como un ecosistema que multiplica valor.')]

doc.add_paragraph().add_run('Preguntas y respuestas').bold=True
table=doc.add_table(rows=1, cols=3); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
set_table_geometry(table,[2550,4350,2460]); set_repeat_table_header(table.rows[0])
headers=['Pregunta','Respuesta analítica','Implicación / recomendación']
for cell,text in zip(table.rows[0].cells,headers):
    shade(cell,'1F4D78'); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; add_para(cell,text,bold=True,color=(255,255,255),size=10.5)
for q,a,imp in questions:
    cells=table.add_row().cells
    for cell in cells: cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    shade(cells[0],'F2F4F7')
    add_para(cells[0],q,bold=True,color=(31,77,120),size=9.7)
    add_para(cells[1],a,size=9.7)
    add_para(cells[2],imp,size=9.7)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
r=p.add_run('Claves para la sustentación oral'); set_font(r,size=13,bold=True,color=(46,116,181))
for text in ['Enfatizar el equilibrio entre crecimiento del parque y diversificación de Val d’Europe.', 'Relacionar la Fase III con gestión de proyectos, gestión de riesgos y creación de valor de largo plazo.', 'Sustentar cada recomendación con los errores iniciales: adaptación cultural insuficiente, demanda sobreestimada y presión financiera.']:
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.1
    r=p.add_run(text); set_font(r,size=10.2)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7)
r=p.add_run('Base documental: Caso Disneyland Paris y diapositiva de preguntas de la Escuela de Postgrado UPC.'); set_font(r,size=8.5,color=(89,89,89),italic=True)

doc.save(OUT)
